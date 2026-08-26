#!/usr/bin/env python3
"""Generate KliqPrint brand documentation as a Word (.docx) file."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


OUT = Path(__file__).resolve().parent / "KliqPrint_Brand_Documents.docx"
ARTIFACT = Path("/opt/cursor/artifacts/KliqPrint_Brand_Documents.docx")


def set_run(run, *, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, bold=True, size=22, color=RGBColor(0x0C, 0x1F, 0x2E))


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, bold=False, size=12, color=RGBColor(0x5C, 0x6B, 0x78))


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=True, size=16, color=RGBColor(0x0C, 0x1F, 0x2E))
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=True, size=13, color=RGBColor(0x1A, 0x33, 0x47))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size=11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def add_page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    add_title(doc, "KliqPrint")
    add_subtitle(doc, "Click · Design · Print · Deliver")
    add_subtitle(doc, "Brand Documents Pack")
    add_body(
        doc,
        "This document contains the official website / business copy for KliqPrint: "
        "About Us, Contact Us, Privacy Policy, Terms of Use, Frequently Asked Questions (FAQs), "
        "Return & Refund Policy, and Sitemap.",
    )
    add_body(
        doc,
        "Replace placeholder fields such as [Your Contact Number], [Your Email Address], "
        "[Your Office / Factory Address], [Your Working Hours], [3/7], and [DD/MM/YYYY] "
        "with final business details before publishing.",
    )

    # ABOUT
    add_page_break(doc)
    add_h1(doc, "About Us")
    add_body(
        doc,
        "KliqPrint is a professional printing, signage, branding and display solutions company "
        "committed to helping businesses create a strong and impactful visual presence. "
        "We provide high-quality customized printing solutions for businesses, institutions, "
        "retailers, event organizers, advertising agencies and individuals.",
    )
    add_body(
        doc,
        "With a focus on quality, customization and reliable service, we combine modern printing "
        "technology with experienced production and finishing to deliver products that meet diverse "
        "branding, promotional and display requirements.",
    )
    add_body(
        doc,
        "Our product range includes custom flags, advertising flags, banners, display stands, "
        "signs & decals, table covers & displays, asset tags, marketing materials and accessories. "
        "From a single customized requirement to bulk commercial orders, our team works closely "
        "with customers from design and artwork preparation through production and final delivery.",
    )
    add_h2(doc, "Why Choose KliqPrint?")
    add_bullets(
        doc,
        [
            "Customized Printing Solutions tailored to your branding and application",
            "Modern Printing Technology for sharp graphics and vibrant colours",
            "Wide Product Range for indoor, outdoor, retail, corporate and event use",
            "Quality Materials & Finishing for a professional appearance",
            "Bulk & Business Orders supported with scalable production",
            "Design & Artwork Assistance to help turn ideas into print-ready products",
            "Reliable Service & Delivery with a focus on customer satisfaction",
        ],
    )
    add_body(
        doc,
        "At KliqPrint, our goal is simple: Click. Design. Print. Deliver. "
        "We make customized printing convenient by bringing design, production and finishing "
        "together under one roof.",
    )
    add_body(
        doc,
        "Whether you need branding for your business, promotional displays for an event, "
        "custom flags, signage or everyday marketing materials, KliqPrint is your partner for "
        "professional print and display solutions.",
    )

    # CONTACT
    add_page_break(doc)
    add_h1(doc, "Contact Us")
    add_body(
        doc,
        "We’d love to hear from you! Whether you need a custom printing quotation, product "
        "information, bulk order pricing, artwork assistance or help choosing the right printing "
        "solution, the KliqPrint team is ready to assist you.",
    )
    add_h2(doc, "Get in Touch")
    add_body(doc, "KliqPrint – Click • Design • Print • Deliver")
    add_bullets(
        doc,
        [
            "Phone / WhatsApp: [Your Contact Number]",
            "Email: [Your Email Address]",
            "Address: [Your Office / Factory Address]",
            "Business Hours: [Monday–Saturday, Your Working Hours]",
        ],
    )
    add_h2(doc, "Custom & Bulk Order Enquiries")
    add_body(
        doc,
        "Looking for custom flags, banners, signage, display stands, table covers, marketing "
        "materials or other customized printing products? Share your required product, size, "
        "quantity, artwork and delivery location, and our team will provide the appropriate "
        "solution and quotation.",
    )
    add_h2(doc, "Send Us an Enquiry")
    add_body(
        doc,
        "Have a question or project in mind? Contact us today and our team will get back to you "
        "as soon as possible.",
    )
    add_body(doc, "KliqPrint — Your Partner for Custom Printing, Signage & Display Solutions.")

    # PRIVACY
    add_page_break(doc)
    add_h1(doc, "Privacy Policy")
    add_body(
        doc,
        "At KliqPrint, we respect your privacy and are committed to protecting the personal "
        "information you share with us. This Privacy Policy explains how we collect, use, store "
        "and protect information when you visit our website, contact us, request a quotation or "
        "purchase our products and services.",
    )
    add_h2(doc, "Information We Collect")
    add_body(
        doc,
        "We may collect information such as your name, company name, phone number, email address, "
        "billing and delivery address, order details, enquiry information and artwork or files "
        "submitted for customized printing.",
    )
    add_body(
        doc,
        "We may also collect certain technical information when you use our website, including "
        "browser type, device information, IP address, pages visited and cookies or similar "
        "technologies.",
    )
    add_h2(doc, "How We Use Your Information")
    add_body(doc, "We may use your information to:")
    add_bullets(
        doc,
        [
            "Process and fulfil orders",
            "Prepare quotations and respond to enquiries",
            "Provide customized printing and design services",
            "Communicate order, payment, production and delivery updates",
            "Provide customer support",
            "Improve our website, products and services",
            "Maintain business and transaction records",
            "Prevent fraud, misuse and security incidents",
            "Send promotional communications where permitted by applicable law",
        ],
    )
    add_h2(doc, "Payment Information")
    add_body(
        doc,
        "Payments may be processed through third-party payment gateways or financial institutions. "
        "KliqPrint does not intentionally store complete card details, UPI PINs, banking passwords "
        "or other sensitive payment authentication information on its own systems.",
    )
    add_h2(doc, "Cookies")
    add_body(
        doc,
        "Our website may use cookies and similar technologies to improve functionality, understand "
        "website usage and enhance your browsing experience. You can control or disable cookies "
        "through your browser settings, although certain website features may not function properly "
        "as a result.",
    )
    add_h2(doc, "Sharing of Information")
    add_body(
        doc,
        "We do not sell or rent your personal information. Information may be shared with trusted "
        "service providers where reasonably necessary to operate our business, including payment "
        "processors, logistics and courier partners, website or technology providers and other "
        "service providers assisting with order fulfilment.",
    )
    add_body(
        doc,
        "We may also disclose information where required by applicable law, regulation, legal "
        "process or government authority.",
    )
    add_h2(doc, "Data Security")
    add_body(
        doc,
        "We take reasonable administrative and technical measures to protect personal information "
        "against unauthorized access, misuse, alteration, disclosure or loss. However, no "
        "internet-based transmission or electronic storage system can be guaranteed to be "
        "completely secure.",
    )
    add_h2(doc, "Data Retention")
    add_body(
        doc,
        "We retain personal information only for as long as reasonably necessary for the purposes "
        "for which it was collected, including order fulfilment, customer support, accounting, "
        "legal and regulatory requirements.",
    )
    add_h2(doc, "Your Rights")
    add_body(
        doc,
        "Subject to applicable law, you may contact us to request access to, correction of or "
        "deletion of your personal information, or to raise questions regarding how your "
        "information is handled.",
    )
    add_h2(doc, "Third-Party Links")
    add_body(
        doc,
        "Our website may contain links to third-party websites or services. KliqPrint is not "
        "responsible for the privacy practices, security or content of those third-party websites.",
    )
    add_h2(doc, "Changes to This Privacy Policy")
    add_body(
        doc,
        "We may update this Privacy Policy periodically to reflect changes in our business "
        "practices, website functionality or applicable legal requirements. The updated version "
        "will be posted on this page with a revised effective date.",
    )
    add_h2(doc, "Contact Us")
    add_body(doc, "For questions or concerns regarding this Privacy Policy or your personal information, please contact:")
    add_body(doc, "KliqPrint")
    add_bullets(
        doc,
        [
            "Email: [Your Email Address]",
            "Phone / WhatsApp: [Your Contact Number]",
            "Address: [Your Registered Business Address]",
        ],
    )
    add_body(doc, "Effective Date: [DD/MM/YYYY]")

    # TERMS
    add_page_break(doc)
    add_h1(doc, "Terms of Use")
    add_body(
        doc,
        "Welcome to KliqPrint. These Terms of Use govern your access to and use of our website, "
        "products and services. By accessing our website, submitting an enquiry or placing an "
        "order, you agree to these terms.",
    )
    add_h2(doc, "Products & Services")
    add_body(
        doc,
        "KliqPrint provides custom printing, flags, banners, signage, display stands, table covers, "
        "marketing materials, accessories and related printing and branding solutions. Product "
        "specifications, colours, sizes, materials and availability may vary depending on the "
        "selected product and customization requirements.",
    )
    add_h2(doc, "Orders & Customization")
    add_body(
        doc,
        "Customers are responsible for providing accurate information regarding size, quantity, "
        "design, text, colours, artwork and other specifications before confirming an order.",
    )
    add_body(
        doc,
        "For customized products, production may begin after approval of the artwork, "
        "specifications and applicable payment. Once customization or production has started, "
        "changes or cancellations may not always be possible.",
    )
    add_h2(doc, "Artwork & Customer-Supplied Content")
    add_body(
        doc,
        "Customers must ensure that they have the necessary rights and permissions to use any "
        "logos, photographs, trademarks, designs, text or other artwork submitted to KliqPrint.",
    )
    add_body(
        doc,
        "By submitting artwork, the customer authorizes KliqPrint to use the supplied material "
        "solely as reasonably necessary to produce the requested products or services.",
    )
    add_h2(doc, "Colours & Product Appearance")
    add_body(
        doc,
        "We make reasonable efforts to reproduce colours and designs accurately. However, slight "
        "variations may occur due to differences in screens, printing processes, materials, inks, "
        "lighting and production batches. Such reasonable variations may not be considered "
        "manufacturing defects.",
    )
    add_h2(doc, "Pricing & Payment")
    add_body(
        doc,
        "Prices may vary depending on product specifications, quantity, customization, taxes, "
        "shipping and other requirements. The applicable price will be communicated or displayed "
        "before order confirmation.",
    )
    add_body(doc, "Orders may require full or partial payment before production or dispatch.")
    add_h2(doc, "Shipping & Delivery")
    add_body(
        doc,
        "Estimated production and delivery timelines are provided in good faith. Delivery times "
        "may vary due to order complexity, courier operations, location, weather, public holidays "
        "or circumstances beyond our reasonable control.",
    )
    add_body(
        doc,
        "Customers are responsible for providing a complete and accurate delivery address and "
        "contact information.",
    )
    add_h2(doc, "Returns, Cancellations & Refunds")
    add_body(
        doc,
        "Returns, cancellations and refunds are subject to our applicable Return, Refund and "
        "Cancellation Policy.",
    )
    add_body(
        doc,
        "Customized, personalized or made-to-order products may not be eligible for return or "
        "cancellation once production has started, except where the product received is defective, "
        "damaged, incorrect or otherwise eligible under applicable law.",
    )
    add_h2(doc, "Intellectual Property")
    add_body(
        doc,
        "Unless otherwise stated, website content belonging to KliqPrint, including its branding, "
        "graphics, product photographs, designs, text and other original materials, is protected "
        "by applicable intellectual property laws.",
    )
    add_body(
        doc,
        "Such content may not be copied, reproduced, distributed or commercially used without "
        "authorization.",
    )
    add_h2(doc, "Acceptable Use")
    add_body(
        doc,
        "You agree not to misuse our website, attempt unauthorized access, interfere with website "
        "functionality, submit fraudulent information or use our services for unlawful purposes.",
    )
    add_h2(doc, "Limitation of Liability")
    add_body(
        doc,
        "To the extent permitted by applicable law, KliqPrint will not be responsible for "
        "indirect or consequential losses arising from the use of our website, products or "
        "services.",
    )
    add_body(
        doc,
        "Nothing in these Terms is intended to exclude or restrict any rights or liabilities that "
        "cannot legally be excluded under applicable law.",
    )
    add_h2(doc, "Third-Party Services")
    add_body(
        doc,
        "Our website or services may use or link to third-party providers, including payment "
        "gateways, courier companies and other service providers. Their respective terms and "
        "policies may apply when you use their services.",
    )
    add_h2(doc, "Changes to These Terms")
    add_body(
        doc,
        "KliqPrint may update these Terms of Use periodically. Any revised terms will become "
        "effective when published on our website unless otherwise stated.",
    )
    add_h2(doc, "Governing Law")
    add_body(
        doc,
        "These Terms will be governed by the laws of India. Any disputes will be subject to the "
        "jurisdiction of the competent courts applicable to KliqPrint’s registered place of "
        "business, subject to applicable consumer protection laws.",
    )
    add_h2(doc, "Contact Us")
    add_body(doc, "For questions regarding these Terms of Use, please contact:")
    add_body(doc, "KliqPrint")
    add_bullets(
        doc,
        [
            "Email: [Your Email Address]",
            "Phone / WhatsApp: [Your Contact Number]",
            "Address: [Your Registered Business Address]",
        ],
    )
    add_body(doc, "Effective Date: [DD/MM/YYYY]")

    # FAQs
    add_page_break(doc)
    add_h1(doc, "Frequently Asked Questions (FAQs)")
    add_body(
        doc,
        "Find answers to some of the most common questions about KliqPrint products, custom "
        "printing, ordering, artwork, payments and delivery.",
    )
    faqs = [
        (
            "1. What products does KliqPrint offer?",
            "KliqPrint provides a wide range of custom printing, signage and display solutions, "
            "including custom flags, advertising flags, banners, display stands, table covers, "
            "signs & decals, asset tags, marketing materials and accessories.",
        ),
        (
            "2. Do you provide customized printing?",
            "Yes. We specialize in customized printing. You can provide your logo, artwork, text, "
            "colours, dimensions and other requirements, and our team will help prepare your "
            "product for production.",
        ),
        (
            "3. Can I order products in custom sizes?",
            "Yes. Custom sizes are available for many products. Availability depends on the "
            "product, material and printing requirements.",
        ),
        (
            "4. Do you accept bulk and corporate orders?",
            "Yes. We accept bulk, corporate, institutional, event and reseller orders. Contact us "
            "with the product, size, quantity and delivery location to request a quotation.",
        ),
        (
            "5. How can I request a quotation?",
            "Send us your requirements through our website, email, phone or WhatsApp. For a faster "
            "quotation, please provide the product name, dimensions, quantity, artwork/design "
            "requirements and delivery location.",
        ),
        (
            "6. What artwork files do you accept?",
            "Depending on the product, we may accept commonly used formats such as PDF, AI, EPS, "
            "CDR, PSD, PNG and high-resolution JPG/JPEG. Vector or high-resolution artwork is "
            "generally recommended for the best printing results.",
        ),
        (
            "7. Can KliqPrint help with artwork or design?",
            "Yes. Our team can assist with basic artwork preparation, sizing and print-ready file "
            "setup, depending on the requirements of your order. Additional design charges may "
            "apply for extensive design work.",
        ),
        (
            "8. Will the printed colours exactly match my screen?",
            "We make reasonable efforts to achieve accurate colour reproduction. However, slight "
            "colour variations can occur because of screen settings, printing methods, inks, "
            "materials and production batches.",
        ),
        (
            "9. How long does production take?",
            "Production time depends on the product, quantity, customization and current workload. "
            "An estimated production and dispatch timeline can be provided when your order is "
            "confirmed.",
        ),
        (
            "10. Do you deliver across India?",
            "Yes, we can ship orders to locations across India, subject to courier serviceability. "
            "Delivery charges and timelines may vary according to the destination, order size and "
            "shipping method.",
        ),
        (
            "11. Do you accept international orders?",
            "International order availability depends on the product, quantity and destination. "
            "Contact our team with your requirements and destination country for assistance with "
            "export orders.",
        ),
        (
            "12. Can I cancel or modify a customized order?",
            "Please contact us as soon as possible. Changes or cancellations may be possible "
            "before production begins. Once a customized or personalized product has entered "
            "production, cancellation or modification may not be possible.",
        ),
        (
            "13. What if I receive a damaged or incorrect product?",
            "Please contact KliqPrint promptly with your order details and clear photographs or "
            "videos of the product and packaging. Our team will review the issue in accordance "
            "with the applicable return, replacement and refund policy.",
        ),
        (
            "14. How can I contact KliqPrint?",
            "For product enquiries, quotations, custom orders or customer support, contact us at:\n"
            "Phone / WhatsApp: [Your Contact Number]\n"
            "Email: [Your Email Address]\n"
            "Address: [Your Business Address]",
        ),
    ]
    for q, a in faqs:
        add_h2(doc, q)
        for para in a.split("\n"):
            add_body(doc, para)
    add_body(doc, "KliqPrint — Click • Design • Print • Deliver")

    # RETURN & REFUND
    add_page_break(doc)
    add_h1(doc, "Return & Refund Policy")
    add_body(
        doc,
        "At KliqPrint, we are committed to delivering quality customized printing and display "
        "products. As many of our products are custom-made, personalized or produced according to "
        "specific customer requirements, returns and refunds are subject to the conditions below.",
    )
    add_h2(doc, "Return Eligibility")
    add_body(doc, "A return or replacement may be considered if:")
    add_bullets(
        doc,
        [
            "The product received is damaged or defective.",
            "You received an incorrect product, size or quantity compared with the confirmed order.",
            "There is a significant printing or manufacturing defect.",
            "The product is materially different from the specifications approved before production.",
        ],
    )
    add_body(doc, "Please contact us within [3/7] days of delivery to report an eligible issue.")
    add_h2(doc, "Customized & Personalized Products")
    add_body(
        doc,
        "Products manufactured according to customer-specific designs, logos, artwork, sizes, "
        "colours, text or other customization requirements are generally non-returnable and "
        "non-refundable unless they are damaged, defective or materially different from the "
        "approved order.",
    )
    add_h2(doc, "How to Request a Return or Replacement")
    add_body(doc, "To report an issue, please contact our customer support team and provide:")
    add_bullets(
        doc,
        [
            "Your order number/invoice number",
            "A description of the issue",
            "Clear photographs of the product",
            "Photographs of the packaging, where relevant",
            "An unboxing video, if available",
        ],
    )
    add_body(
        doc,
        "Our team will review the information and advise you regarding the appropriate "
        "replacement, return or refund, where applicable.",
    )
    add_h2(doc, "Non-Returnable Situations")
    add_body(
        doc,
        "Returns may not be accepted for products that have been used, installed, altered, washed "
        "or damaged after delivery, products ordered with incorrect specifications supplied by the "
        "customer, minor colour variations caused by screens/materials/printing processes, or "
        "customized products produced according to approved artwork and specifications.",
    )
    add_h2(doc, "Order Cancellation")
    add_body(
        doc,
        "If you need to cancel an order, please contact us immediately. Orders may be cancelled "
        "before production begins, subject to the status of the order.",
    )
    add_body(
        doc,
        "Once printing, customization or production has started, cancellation may not be possible.",
    )
    add_h2(doc, "Refunds")
    add_body(
        doc,
        "If a refund is approved, it will generally be processed to the original payment method "
        "or another mutually agreed method. Processing time may vary depending on the payment "
        "provider or financial institution.",
    )
    add_body(
        doc,
        "Any applicable shipping, design, customization or production charges may be deducted "
        "where permitted and appropriate.",
    )
    add_h2(doc, "Replacement")
    add_body(
        doc,
        "Where an eligible manufacturing defect, damage or incorrect product is confirmed, "
        "KliqPrint may offer a replacement instead of a refund where appropriate.",
    )
    add_h2(doc, "Shipping Damage")
    add_body(
        doc,
        "If your package arrives visibly damaged, please take clear photographs of the outer "
        "packaging and product before disposing of the packaging. Contact us promptly so we can "
        "review the issue.",
    )
    add_h2(doc, "Contact Us")
    add_body(doc, "For return, replacement, cancellation or refund assistance, please contact:")
    add_body(doc, "KliqPrint")
    add_bullets(
        doc,
        [
            "Phone / WhatsApp: [Your Contact Number]",
            "Email: [Your Email Address]",
            "Address: [Your Registered Business Address]",
        ],
    )
    add_body(doc, "Effective Date: [DD/MM/YYYY]")

    # SITEMAP
    add_page_break(doc)
    add_h1(doc, "Sitemap")
    add_body(doc, "Quickly explore KliqPrint and find the products, services and information you need.")
    add_h2(doc, "Main Pages")
    add_bullets(doc, ["Home", "About Us", "Contact Us", "FAQs"])
    add_h2(doc, "Products")
    add_bullets(
        doc,
        [
            "Flags",
            "Custom Flags",
            "Advertising Flags",
            "Banners",
            "Display Stands",
            "Signs & Decals",
            "Table Covers & Displays",
            "Asset Tags",
            "Marketing Materials",
            "Accessories",
        ],
    )
    add_h2(doc, "Custom Printing Services")
    add_bullets(
        doc,
        [
            "Custom Design & Printing",
            "Business Branding Solutions",
            "Event & Promotional Printing",
            "Corporate & Bulk Orders",
            "Custom Size Printing",
            "Artwork & Design Assistance",
        ],
    )
    add_h2(doc, "Customer Support")
    add_bullets(
        doc,
        [
            "Contact Us",
            "Request a Quote",
            "FAQs",
            "Order & Delivery Information",
            "Return & Refund Policy",
        ],
    )
    add_h2(doc, "Legal & Policies")
    add_bullets(doc, ["Privacy Policy", "Terms of Use", "Return & Refund Policy"])
    add_h2(doc, "Connect With KliqPrint")
    add_bullets(doc, ["Phone / WhatsApp", "Email", "Social Media"])
    add_body(doc, "KliqPrint — Click • Design • Print • Deliver")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    ARTIFACT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    doc.save(ARTIFACT)
    print(f"Wrote {OUT}")
    print(f"Wrote {ARTIFACT}")


if __name__ == "__main__":
    build()
